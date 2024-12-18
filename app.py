import streamlit as st
from anthropic import Anthropic
import pypdf
import docx
import io
import pandas as pd
from difflib import SequenceMatcher
import pdfplumber
import re
from typing import List, Dict
import fitz  # PyMuPDF

# Streamlit page config
st.set_page_config(page_title="Sales Quote Comparison", page_icon="💼", layout="wide")

# Initialize Streamlit page
st.title("Sales Quote Line Item Comparison")
st.markdown("Using Claude 3 Opus for Enhanced Analysis")

# Secure API key handling
try:
    if 'anthropic_client' not in st.session_state:
        anthropic_api_key = st.secrets["ANTHROPIC_API_KEY"]
        st.session_state.anthropic_client = Anthropic(api_key=anthropic_api_key)
except Exception as e:
    st.error("Please configure your API key in Streamlit secrets")
    st.stop()

# Initialize session state
if 'quote1' not in st.session_state:
    st.session_state.quote1 = None
if 'quote2' not in st.session_state:
    st.session_state.quote2 = None
if 'comparison_results' not in st.session_state:
    st.session_state.comparison_results = None

def extract_pdf_text_pdfplumber(file) -> str:
    """Extract text from PDF using pdfplumber with enhanced formatting."""
    try:
        with pdfplumber.open(file) as pdf:
            pages_text = []
            for page in pdf.pages:
                # Extract text with better handling of tables and layouts
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                # Extract tables and format them properly
                tables = page.extract_tables()
                formatted_tables = []
                for table in tables:
                    formatted_table = '\n'.join(['\t'.join([str(cell) if cell else '' for cell in row]) for row in table])
                    formatted_tables.append(formatted_table)
                
                # Combine regular text and tables
                page_text = text + '\n' + '\n'.join(formatted_tables)
                pages_text.append(page_text)
            
            return '\n\n'.join(pages_text)
    except Exception as e:
        st.error(f"Error in pdfplumber extraction: {str(e)}")
        return ""

def extract_pdf_text_pymupdf(file) -> str:
    """Extract text from PDF using PyMuPDF for better formatting."""
    try:
        file_bytes = file.getvalue()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        pages_text = []
        for page in doc:
            blocks = page.get_text("dict")["blocks"]
            page_text = []
            
            for block in blocks:
                if block["type"] == 0:  # Text block
                    for line in block["lines"]:
                        line_text = ""
                        for span in line["spans"]:
                            line_text += span["text"]
                        page_text.append(line_text)
                elif block["type"] == 1:  # Image block
                    page_text.append("[Image]")
            
            pages_text.append("\n".join(page_text))
        
        doc.close()
        return "\n\n".join(pages_text)
    except Exception as e:
        st.error(f"Error in PyMuPDF extraction: {str(e)}")
        return ""

def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s.,;:$%()-]', '', text)
    text = text.replace('\r', '\n')
    text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
    return text

def extract_tables_from_text(text: str) -> List[List[str]]:
    """Attempt to identify and extract tabular data from text."""
    tables = []
    current_table = []
    
    lines = text.split('\n')
    for line in lines:
        if '\t' in line or '  ' in line:
            cells = re.split(r'\t|  +', line.strip())
            current_table.append(cells)
        elif current_table:
            if len(current_table) > 1:
                tables.append(current_table)
            current_table = []
    
    if current_table and len(current_table) > 1:
        tables.append(current_table)
    
    return tables

def read_file(file) -> Dict[str, str]:
    """Enhanced file reading with better PDF handling and text preprocessing."""
    try:
        if file.type == "application/pdf":
            text_pdfplumber = extract_pdf_text_pdfplumber(file)
            text_pymupdf = extract_pdf_text_pymupdf(file)
            
            text = text_pdfplumber if len(text_pdfplumber) > len(text_pymupdf) else text_pymupdf
            tables = extract_tables_from_text(text)
            formatted_tables = ['\n'.join(['\t'.join(row) for row in table]) for table in tables]
            final_text = text + '\n\n' + '\n\n'.join(formatted_tables)
            final_text = clean_text(final_text)
            
            return {
                'text': final_text,
                'tables': tables,
                'raw_text': text
            }
            
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file)
            text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            tables = [[cell.text for cell in row.cells] for table in doc.tables for row in table.rows]
            return {
                'text': clean_text(text),
                'tables': tables,
                'raw_text': text
            }
        else:  # txt files
            text = file.getvalue().decode('utf-8')
            return {
                'text': clean_text(text),
                'tables': extract_tables_from_text(text),
                'raw_text': text
            }
    except Exception as e:
        st.error(f"Error reading file {file.name}: {str(e)}")
        return {'text': '', 'tables': [], 'raw_text': ''}
        # Create two columns for file uploads
col1, col2 = st.columns(2)

# Modified file upload handling
with col1:
    st.subheader("Quote 1")
    quote1_file = st.file_uploader(
        "Upload first quote",
        type=['pdf', 'docx', 'txt'],
        key="quote1_uploader"
    )
    if quote1_file:
        extracted_data = read_file(quote1_file)
        st.session_state.quote1 = {
            'name': quote1_file.name,
            'content': extracted_data['text'],
            'tables': extracted_data['tables'],
            'raw_content': extracted_data['raw_text']
        }
        st.success(f"Quote 1 uploaded: {quote1_file.name}")
        
        # Show preview with tabs for different views
        preview_tab1, preview_tab2 = st.tabs(["Processed Text", "Raw Text"])
        with preview_tab1:
            st.text(extracted_data['text'][:1000] + "...")
        with preview_tab2:
            st.text(extracted_data['raw_text'][:1000] + "...")
            
        # Show detected tables
        if extracted_data['tables']:
            with st.expander("View Detected Tables"):
                for i, table in enumerate(extracted_data['tables']):
                    st.markdown(f"**Table {i+1}**")
                    st.dataframe(pd.DataFrame(table))

with col2:
    st.subheader("Quote 2")
    quote2_file = st.file_uploader(
        "Upload second quote",
        type=['pdf', 'docx', 'txt'],
        key="quote2_uploader"
    )
    if quote2_file:
        extracted_data = read_file(quote2_file)
        st.session_state.quote2 = {
            'name': quote2_file.name,
            'content': extracted_data['text'],
            'tables': extracted_data['tables'],
            'raw_content': extracted_data['raw_text']
        }
        st.success(f"Quote 2 uploaded: {quote2_file.name}")
        
        # Show preview with tabs for different views
        preview_tab1, preview_tab2 = st.tabs(["Processed Text", "Raw Text"])
        with preview_tab1:
            st.text(extracted_data['text'][:1000] + "...")
        with preview_tab2:
            st.text(extracted_data['raw_text'][:1000] + "...")
            
        # Show detected tables
        if extracted_data['tables']:
            with st.expander("View Detected Tables"):
                for i, table in enumerate(extracted_data['tables']):
                    st.markdown(f"**Table {i+1}**")
                    st.dataframe(pd.DataFrame(table))

# Compare button
if st.button("Compare Quotes") and st.session_state.quote1 and st.session_state.quote2:
    with st.spinner("Performing detailed line item analysis with Claude 3 Opus..."):
        try:
            # Split the analysis into smaller chunks if needed
            extraction_prompt = f"""
            You are a technical expert in analyzing and comparing detailed quotes. 
            I need you to do a thorough line-by-line comparison of these two quotes.
            
            IMPORTANT INSTRUCTIONS:
            1. Analyze EVERY SINGLE LINE ITEM in detail
            2. Do not summarize or skip any items
            3. Look for exact and partial matches even if wording is different
            4. Pay special attention to technical specifications and model numbers
            5. Identify items that might be the same but described differently
            
            Quote 1 ({st.session_state.quote1['name']}):
            {st.session_state.quote1['content']}

            Quote 2 ({st.session_state.quote2['name']}):
            {st.session_state.quote2['content']}

            Create a DETAILED comparison table with EXACTLY these columns:
            | Line Item | Quote 1 Description | Quote 2 Description | Match Status | Notes |

            Match Status Rules:
            - ✓ = Exact match (identical specifications)
            - ~ = Partial match (similar but with differences)
            - [1] = Only in Quote 1
            - [2] = Only in Quote 2

            IMPORTANT:
            - Include EVERY line item from both quotes
            - Use EXACT descriptions from each quote
            - Note ANY differences in specifications
            - If items are similar but not identical, mark as partial match
            - Be extremely thorough in comparison
            """
            
            # Initial analysis with correct max_tokens
            initial_response = st.session_state.anthropic_client.messages.create(
                model="claude-3-opus-20240229",
                max_tokens=4096,  # Corrected to maximum allowed
                messages=[{
                    "role": "user",
                    "content": extraction_prompt
                }]
            )
            
            # Create tabs for different views
            tab1, tab2, tab3 = st.tabs([
                "Line Item Comparison", 
                "Detailed Analysis", 
                "Summary"
            ])
            
            with tab1:
                st.markdown("### Complete Line Item Comparison")
                
                # Enhanced styling for better visibility
                st.markdown("""
                <style>
                .exact-match { color: #00aa00; font-weight: bold; background-color: #f0fff0; }
                .partial-match { color: #ff6600; font-weight: bold; background-color: #fff6f0; }
                .unique-item { color: #0066ff; font-weight: bold; background-color: #f0f6ff; }
                .different { background-color: #fff3f3; }
                th { background-color: #f0f2f6; }
                td { min-width: 150px; }
                </style>
                """, unsafe_allow_html=True)
                
                # Visualization prompt
                viz_response = st.session_state.anthropic_client.messages.create(
                    model="claude-3-opus-20240229",
                    max_tokens=4096,
                    messages=[
                        {"role": "assistant", "content": initial_response.content[0].text},
                        {"role": "user", "content": "Format the above analysis as a clear markdown table showing all items and their comparisons."}
                    ]
                )
                
                # Display the full comparison table
                st.markdown(viz_response.content[0].text)
                
                # Download options
                st.download_button(
                    "Download Full Comparison as CSV",
                    viz_response.content[0].text,
                    "detailed_line_item_comparison.csv",
                    "text/csv"
                )
            
            with tab2:
                st.markdown("### Detailed Analysis")
                
                analysis_response = st.session_state.anthropic_client.messages.create(
                    model="claude-3-opus-20240229",
                    max_tokens=4096,
                    messages=[
                        {"role": "assistant", "content": initial_response.content[0].text},
                        {"role": "user", "content": "Provide a detailed analysis of all differences found, including partial matches and unique items."}
                    ]
                )
                
                st.markdown(analysis_response.content[0].text)
            
            with tab3:
                st.markdown("### Comparison Summary")
                
                summary_response = st.session_state.anthropic_client.messages.create(
                    model="claude-3-opus-20240229",
                    max_tokens=4096,
                    messages=[
                        {"role": "assistant", "content": initial_response.content[0].text},
                        {"role": "user", "content": "Provide a statistical summary of the comparison, including counts of exact matches, partial matches, and unique items."}
                    ]
                )
                
                st.markdown(summary_response.content[0].text)
                
                # Metrics display
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Total Items", "Calculated")
                with col2:
                    st.metric("Exact Matches", "Calculated")
                with col3:
                    st.metric("Partial Matches", "Calculated")
                with col4:
                    st.metric("Unique Items", "Calculated")

        except Exception as e:
            st.error(f"Error in comparison: {str(e)}")
            st.error("Full error details:", str(e))
